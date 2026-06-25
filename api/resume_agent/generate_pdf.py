#!/usr/bin/env python3
"""
Professional PDF + DOCX resume builder.
"""

import io
import re
import sys
import json

# ΓöÇΓöÇ Unicode sanitizer ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ

def sanitize(text: str) -> str:
    if not text:
        return ''
    text = str(text)
    table = {
        'ΓÇô': '-', 'ΓÇö': '-', 'ΓÇò': '-',
        'ΓÇÿ': "'", 'ΓÇÖ': "'", 'ΓÇÜ': "'",
        'ΓÇ£': '"', 'ΓÇ¥': '"', 'ΓÇ₧': '"',
        'ΓÇó': '*', '┬╖': '*', 'Γû¬': '*',
        'ΓÇª': '...', '┬á': ' ', '┬«': '(R)',
        'Γäó': '(TM)', '┬╢': '', '┬⌐': '(C)',
        '\xc3\xa2': '-', '\xe2\x80\x93': '-', '\xe2\x80\x94': '-',
    }
    for src, dst in table.items():
        text = text.replace(src, dst)
    # strip anything outside printable ASCII / latin-1 range
    text = re.sub(r'[^\x20-\x7E\xA0-\xFF]', '', text)
    # encode/decode round-trip to catch any remaining issues
    text = text.encode('latin-1', errors='replace').decode('latin-1')
    return text


def s(v) -> str:
    return sanitize(v) if v else ''


# ΓöÇΓöÇ PDF builder ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ

def build_pdf_bytes(data: dict) -> bytes:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_RIGHT
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer,
        HRFlowable, Table, TableStyle, KeepTogether,
    )

    DARK  = colors.HexColor('#1a3560')
    MID   = colors.HexColor('#2c5282')
    GREY  = colors.HexColor('#555555')
    BLACK = colors.black
    WHITE = colors.white

    # Create all styles as standalone objects ΓÇö avoids the global singleton race condition
    def st(**kw):
        return ParagraphStyle('_', **kw)

    ST_NAME    = st(fontName='Helvetica-Bold', fontSize=22,
                    textColor=WHITE, alignment=TA_CENTER, leading=26)
    ST_CONTACT = st(fontName='Helvetica', fontSize=9,
                    textColor=colors.HexColor('#d0ddf5'), alignment=TA_CENTER, leading=13)
    ST_SEC     = st(fontName='Helvetica-Bold', fontSize=10,
                    textColor=DARK, spaceBefore=8, spaceAfter=1, leading=13)
    ST_BODY    = st(fontName='Helvetica', fontSize=9.5,
                    textColor=BLACK, leading=14, spaceAfter=2)
    ST_BULLET  = st(fontName='Helvetica', fontSize=9,
                    textColor=colors.HexColor('#222222'), leading=13,
                    leftIndent=10, spaceAfter=1)
    ST_JOBT    = st(fontName='Helvetica-Bold', fontSize=10,
                    textColor=DARK, spaceBefore=5, spaceAfter=0, leading=13)
    ST_JOBM    = st(fontName='Helvetica', fontSize=8.5,
                    textColor=GREY, spaceAfter=2, leading=12)
    ST_JOBDATE = st(fontName='Helvetica-Oblique', fontSize=8.5,
                    textColor=GREY, spaceAfter=2, leading=12, alignment=TA_RIGHT)
    ST_EDUT    = st(fontName='Helvetica-Bold', fontSize=9.5,
                    textColor=DARK, spaceBefore=4, spaceAfter=0, leading=13)
    ST_EDUS    = st(fontName='Helvetica', fontSize=8.5,
                    textColor=GREY, spaceAfter=2, leading=12)
    ST_EDUDATE = st(fontName='Helvetica-Oblique', fontSize=8.5,
                    textColor=GREY, spaceAfter=2, leading=12, alignment=TA_RIGHT)
    ST_CERT    = st(fontName='Helvetica', fontSize=9,
                    textColor=BLACK, leading=13, spaceAfter=1)
    ST_LANG    = st(fontName='Helvetica', fontSize=9,
                    textColor=BLACK, leading=13, spaceAfter=1)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=letter,
        leftMargin=0.65*inch, rightMargin=0.65*inch,
        topMargin=0.55*inch, bottomMargin=0.55*inch,
    )

    W = letter[0] - 0.65*inch*2

    ROW_PAD = [
        ('LEFTPADDING',  (0,0), (-1,-1), 0),
        ('RIGHTPADDING', (0,0), (-1,-1), 0),
        ('TOPPADDING',   (0,0), (-1,-1), 0),
        ('BOTTOMPADDING',(0,0), (-1,-1), 0),
        ('VALIGN',       (0,0), (-1,-1), 'BOTTOM'),
    ]

    def hr():
        return HRFlowable(width='100%', thickness=0.75, color=MID,
                          spaceAfter=4, spaceBefore=1)

    def section(title):
        return [Paragraph(title.upper(), ST_SEC), hr()]

    def two_col(left_para, right_text, right_style):
        t = Table([[left_para, Paragraph(right_text, right_style)]],
                  colWidths=[W*0.7, W*0.3])
        t.setStyle(TableStyle(ROW_PAD))
        return t

    story = []

    # ΓöÇΓöÇ Header ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
    contact_bits = [s(p) for p in [
        data.get('email'), data.get('phone'),
        data.get('location'), data.get('linkedin'),
    ] if p]
    hdr = Table(
        [[Paragraph(s(data.get('name', 'Resume')), ST_NAME)],
         [Paragraph('   |   '.join(contact_bits), ST_CONTACT)]],
        colWidths=[W],
    )
    hdr.setStyle(TableStyle([
        ('BACKGROUND',    (0,0), (-1,-1), DARK),
        ('TOPPADDING',    (0,0), (-1,-1), 12),
        ('BOTTOMPADDING', (0,0), (-1,-1), 10),
        ('LEFTPADDING',   (0,0), (-1,-1), 8),
        ('RIGHTPADDING',  (0,0), (-1,-1), 8),
    ]))
    story.append(hdr)
    story.append(Spacer(1, 8))

    # ΓöÇΓöÇ Summary ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
    if data.get('summary'):
        story += section('Professional Summary')
        story.append(Paragraph(s(data['summary']), ST_BODY))

    # ΓöÇΓöÇ Skills ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
    skills = data.get('skills') or []
    if skills:
        story += section('Core Skills')
        ST_SKILL = st(fontName='Helvetica', fontSize=9, textColor=BLACK,
                      leading=13, spaceAfter=1)
        # Detect if skills are long category strings (contain ':') ΓåÆ 1-col list
        # Otherwise lay out as a 2-column table so text wraps properly
        long_form = any(':' in str(sk) and len(str(sk)) > 40 for sk in skills)
        if long_form:
            for sk in skills:
                txt = s(sk)
                if txt:
                    story.append(Paragraph(f'-  {txt}', ST_SKILL))
        else:
            rows, row = [], []
            for sk in skills:
                row.append(Paragraph(s(sk), ST_SKILL))
                if len(row) == 2:
                    rows.append(row); row = []
            if row:
                rows.append(row + [Paragraph('', ST_SKILL)] * (2 - len(row)))
            if rows:
                t = Table(rows, colWidths=[W/2, W/2])
                t.setStyle(TableStyle([
                    ('TOPPADDING',    (0,0), (-1,-1), 1),
                    ('BOTTOMPADDING', (0,0), (-1,-1), 1),
                    ('LEFTPADDING',   (0,0), (-1,-1), 2),
                    ('RIGHTPADDING',  (0,0), (-1,-1), 2),
                    ('VALIGN',        (0,0), (-1,-1), 'TOP'),
                ]))
                story.append(t)
        story.append(Spacer(1, 4))

    # ΓöÇΓöÇ Experience ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
    exp = data.get('experience') or []
    if exp:
        story += section('Experience')
        for job in exp:
            title   = s(job.get('title', ''))
            company = s(job.get('company', ''))
            dates   = s(job.get('dates', ''))
            bullets = job.get('bullets') or []

            block = [
                two_col(Paragraph(title, ST_JOBT), dates, ST_JOBDATE),
                Paragraph(company, ST_JOBM),
            ]
            for b in bullets:
                txt = s(b)
                if txt:
                    block.append(Paragraph(f'-  {txt}', ST_BULLET))
            story.append(KeepTogether(block))

    # ΓöÇΓöÇ Education ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
    edu = data.get('education') or []
    if edu:
        story += section('Education')
        for e in edu:
            degree  = s(e.get('degree', ''))
            school  = s(e.get('school', ''))
            dates   = s(e.get('dates', ''))
            details = s(e.get('details', ''))
            sub = f'{school}  |  {details}' if school and details else school or details
            block = [two_col(Paragraph(degree, ST_EDUT), dates, ST_EDUDATE)]
            if sub:
                block.append(Paragraph(sub, ST_EDUS))
            story.append(KeepTogether(block))

    # ΓöÇΓöÇ Certifications ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
    certs = data.get('certifications') or []
    if certs:
        story += section('Certifications')
        for c in certs:
            txt = s(c)
            if txt:
                story.append(Paragraph(f'-  {txt}', ST_CERT))

    # ΓöÇΓöÇ Languages ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
    langs = data.get('languages') or []
    if langs:
        story += section('Languages')
        for lang in langs:
            txt = s(lang)
            if txt:
                story.append(Paragraph(f'-  {txt}', ST_LANG))

    doc.build(story)
    buf.seek(0)
    return buf.read()


# ΓöÇΓöÇ DOCX builder ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ

def build_docx_bytes(data: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    DARK_R, DARK_G, DARK_B = 26, 53, 96   # #1a3560

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin   = Inches(0.75)
        section.right_margin  = Inches(0.75)

    def set_font(run, size, bold=False, italic=False,
                 color=None, name='Calibri'):
        run.font.name  = name
        run.font.size  = Pt(size)
        run.font.bold  = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = RGBColor(*color)

    def add_heading_para(text, size=11, color=(26, 53, 96)):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(text.upper())
        set_font(run, size, bold=True, color=color)
        return p

    def add_hr(para):
        """Add a bottom border to a paragraph (acts as a rule under section head)."""
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '2c5282')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def section_head(title):
        p = add_heading_para(title)
        add_hr(p)

    # ΓöÇΓöÇ Header ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(2)
    nr = name_p.add_run(s(data.get('name', '')))
    set_font(nr, 22, bold=True, color=(DARK_R, DARK_G, DARK_B))

    contact_bits = [s(p) for p in [
        data.get('email'), data.get('phone'),
        data.get('location'), data.get('linkedin'),
    ] if p]
    if contact_bits:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(6)
        cr = cp.add_run('   |   '.join(contact_bits))
        set_font(cr, 9, color=(80, 80, 80))

    # ΓöÇΓöÇ Summary ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
    if data.get('summary'):
        section_head('Professional Summary')
        p = doc.add_paragraph(s(data['summary']))
        p.style.font.size = Pt(9.5)
        p.paragraph_format.space_after = Pt(4)

    # ΓöÇΓöÇ Skills ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
    skills = data.get('skills') or []
    if skills:
        section_head('Core Skills')
        long_form = any(':' in str(sk) and len(str(sk)) > 40 for sk in skills)
        if long_form:
            # Category-style long strings ΓÇö one per bullet line
            for sk in skills:
                txt = s(sk)
                if txt:
                    bp = doc.add_paragraph(style='List Bullet')
                    bp.paragraph_format.space_after = Pt(1)
                    br = bp.add_run(txt)
                    set_font(br, 9)
        else:
            # Short individual skills ΓÇö 2-column table
            rows_data, row = [], []
            for sk in skills:
                row.append(s(sk))
                if len(row) == 2:
                    rows_data.append(row); row = []
            if row:
                row += [''] * (2 - len(row))
                rows_data.append(row)
            tbl = doc.add_table(rows=len(rows_data), cols=2)
            tbl.style = 'Table Grid'
            for r_idx, row_data in enumerate(rows_data):
                for c_idx, cell_text in enumerate(row_data):
                    cell = tbl.rows[r_idx].cells[c_idx]
                    cell.text = cell_text
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(9)
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcBorders = OxmlElement('w:tcBorders')
                    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                        el = OxmlElement(f'w:{side}')
                        el.set(qn('w:val'), 'none')
                        tcBorders.append(el)
                    tcPr.append(tcBorders)

    # ΓöÇΓöÇ Experience ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
    exp = data.get('experience') or []
    if exp:
        section_head('Experience')
        for job in exp:
            title   = s(job.get('title', ''))
            company = s(job.get('company', ''))
            dates   = s(job.get('dates', ''))
            bullets = job.get('bullets') or []

            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after  = Pt(0)
            tr = p.add_run(title)
            set_font(tr, 10, bold=True, color=(DARK_R, DARK_G, DARK_B))
            if dates:
                p.add_run('   ')
                dr = p.add_run(dates)
                set_font(dr, 8.5, italic=True, color=(100, 100, 100))

            if company:
                cp2 = doc.add_paragraph(company)
                cp2.paragraph_format.space_after = Pt(2)
                for run in cp2.runs:
                    set_font(run, 8.5, color=(80, 80, 80))

            for b in bullets:
                txt = s(b)
                if txt:
                    bp = doc.add_paragraph(style='List Bullet')
                    bp.paragraph_format.left_indent  = Inches(0.2)
                    bp.paragraph_format.space_after  = Pt(1)
                    br = bp.add_run(txt)
                    set_font(br, 9)

    # ΓöÇΓöÇ Education ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
    edu = data.get('education') or []
    if edu:
        section_head('Education')
        for e in edu:
            degree  = s(e.get('degree', ''))
            school  = s(e.get('school', ''))
            dates   = s(e.get('dates', ''))
            details = s(e.get('details', ''))

            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(0)
            dr2 = p.add_run(degree)
            set_font(dr2, 10, bold=True, color=(DARK_R, DARK_G, DARK_B))
            if dates:
                p.add_run('   ')
                dtr = p.add_run(dates)
                set_font(dtr, 8.5, italic=True, color=(100, 100, 100))

            sub = school
            if details:
                sub = f'{school}  |  {details}' if school else details
            if sub:
                sp = doc.add_paragraph(sub)
                sp.paragraph_format.space_after = Pt(2)
                for run in sp.runs:
                    set_font(run, 8.5, color=(80, 80, 80))

    # ΓöÇΓöÇ Certifications ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
    certs = data.get('certifications') or []
    if certs:
        section_head('Certifications')
        for c in certs:
            txt = s(c)
            if txt:
                cp = doc.add_paragraph(style='List Bullet')
                cp.paragraph_format.space_after = Pt(1)
                cr2 = cp.add_run(txt)
                set_font(cr2, 9)

    # ΓöÇΓöÇ Languages ΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇΓöÇ
    langs = data.get('languages') or []
    if langs:
        section_head('Languages')
        for lang in langs:
            txt = s(lang)
            if txt:
                lp = doc.add_paragraph(style='List Bullet')
                lp.paragraph_format.space_after = Pt(1)
                lr = lp.add_run(txt)
                set_font(lr, 9)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print('Usage: python generate_pdf.py <resume.json> <output.pdf>', file=sys.stderr)
        sys.exit(1)
    with open(sys.argv[1]) as f:
        resume_data = json.load(f)
    out = sys.argv[2]
    if out.endswith('.docx'):
        with open(out, 'wb') as f:
            f.write(build_docx_bytes(resume_data))
        print(f'DOCX saved: {out}')
    else:
        with open(out, 'wb') as f:
            f.write(build_pdf_bytes(resume_data))
        print(f'PDF saved: {out}')
